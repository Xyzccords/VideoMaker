"""
traducir_core.py
-----------------
Version parametrizada de traducir_novelas.py, pensada para ser llamada desde
una interfaz grafica (interfaz.py) en vez de ejecutarse por linea de comandos.

Expone la funcion traducir_carpeta(cfg, log, progreso, should_stop) que hace
lo mismo que el script original: recorre una carpeta buscando .docx, los
traduce respetando el formato basico, y guarda el resultado en una carpeta
hermana "(traducido)", saltando los capitulos que ya esten traducidos.
"""

import time
from pathlib import Path

from docx import Document
from deep_translator import GoogleTranslator


LIMITE_CARACTERES = 4500
PAUSA_ENTRE_LLAMADAS = 0.3
MAX_REINTENTOS = 4
SEPARADOR = "@@@"


def _traducir_llamada(translator, texto, log, max_reintentos=MAX_REINTENTOS):
    for intento in range(1, max_reintentos + 1):
        try:
            traducido = translator.translate(texto)
            return traducido if traducido else texto
        except Exception as e:
            if intento == max_reintentos:
                log(f"  [aviso] no se pudo traducir un fragmento, se deja en original: {e}")
                return texto
            time.sleep(1.5 * intento)
    return texto


def _partir_texto(texto, limite=LIMITE_CARACTERES):
    if len(texto) <= limite:
        return [texto]

    partes = []
    actual = ""
    for fragmento in texto.replace("\n", " \n").split(". "):
        candidato = (actual + ". " + fragmento) if actual else fragmento
        if len(candidato) > limite:
            if actual:
                partes.append(actual)
            actual = fragmento
        else:
            actual = candidato
    if actual:
        partes.append(actual)
    return partes


def _traducir_texto(translator, texto, log):
    if not texto or not texto.strip():
        return texto
    trozos = _partir_texto(texto)
    resultado = []
    for trozo in trozos:
        resultado.append(_traducir_llamada(translator, trozo, log))
        time.sleep(PAUSA_ENTRE_LLAMADAS)
    return " ".join(resultado)


def _traducir_lote(translator, textos, log):
    resultados = [None] * len(textos)
    indices = [i for i, t in enumerate(textos) if t and t.strip()]

    i = 0
    while i < len(indices):
        idx = indices[i]
        texto = textos[idx]

        if len(texto) > LIMITE_CARACTERES:
            resultados[idx] = _traducir_texto(translator, texto, log)
            i += 1
            continue

        lote_indices = [idx]
        lote_textos = [texto]
        longitud = len(texto)
        i += 1
        while i < len(indices):
            sig_idx = indices[i]
            sig_texto = textos[sig_idx]
            if len(sig_texto) > LIMITE_CARACTERES:
                break
            extra = len(SEPARADOR) + len(sig_texto)
            if longitud + extra > LIMITE_CARACTERES:
                break
            lote_textos.append(sig_texto)
            lote_indices.append(sig_idx)
            longitud += extra
            i += 1

        if len(lote_textos) == 1:
            resultados[lote_indices[0]] = _traducir_llamada(translator, lote_textos[0], log)
        else:
            combinado = f"\n{SEPARADOR}\n".join(lote_textos)
            traducido = _traducir_llamada(translator, combinado, log)
            partes = [p.strip() for p in traducido.split(SEPARADOR)]

            if len(partes) != len(lote_textos):
                for li, lt in zip(lote_indices, lote_textos):
                    resultados[li] = _traducir_llamada(translator, lt, log)
                    time.sleep(PAUSA_ENTRE_LLAMADAS)
            else:
                for li, parte in zip(lote_indices, partes):
                    resultados[li] = parte

        time.sleep(PAUSA_ENTRE_LLAMADAS)

    for j, t in enumerate(textos):
        if resultados[j] is None:
            resultados[j] = t
    return resultados


def _recolectar_parrafos(doc):
    parrafos = list(doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                parrafos.extend(celda.paragraphs)
    for seccion in doc.sections:
        parrafos.extend(seccion.header.paragraphs)
        parrafos.extend(seccion.footer.paragraphs)
    return parrafos


def _aplicar_texto(paragraph, texto_traducido):
    if not paragraph.runs:
        return
    run_modelo = paragraph.runs[0]
    for run in list(paragraph.runs):
        run.text = ""
    run_modelo.text = texto_traducido


def _traducir_docx(translator, ruta_entrada: Path, ruta_salida: Path, log):
    doc = Document(str(ruta_entrada))
    parrafos = _recolectar_parrafos(doc)
    textos_originales = [p.text for p in parrafos]
    textos_traducidos = _traducir_lote(translator, textos_originales, log)

    for parrafo, texto_traducido in zip(parrafos, textos_traducidos):
        if parrafo.text and parrafo.text.strip():
            _aplicar_texto(parrafo, texto_traducido)

    ruta_salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ruta_salida))


def traducir_carpeta(cfg, log=print, progreso=None, should_stop=None):
    """
    cfg debe traer: carpeta_base (str o Path), idioma_origen, idioma_destino.
    progreso: funcion(hecho, total) llamada tras cada capitulo.
    should_stop: funcion sin argumentos que devuelve True si hay que cancelar.
    """
    should_stop = should_stop or (lambda: False)
    progreso = progreso or (lambda hecho, total: None)

    carpeta_base = Path(cfg["carpeta_base"]).expanduser().resolve()
    if not carpeta_base.is_dir():
        log(f"No se encontro la carpeta: {carpeta_base}")
        return

    carpeta_salida_base = carpeta_base.parent / f"{carpeta_base.name} (traducido)"

    archivos = sorted(carpeta_base.rglob("*.docx"))
    if not archivos:
        log("No se encontraron archivos .docx dentro de esa carpeta.")
        return

    log(f"Se encontraron {len(archivos)} capitulos (.docx).")
    log(f"Los archivos traducidos se guardaran en:\n  {carpeta_salida_base}\n")

    pendientes = []
    for archivo in archivos:
        relativo = archivo.relative_to(carpeta_base)
        destino = carpeta_salida_base / relativo
        if destino.exists():
            continue
        pendientes.append((archivo, destino))

    ya_hechos = len(archivos) - len(pendientes)
    if ya_hechos:
        log(f"{ya_hechos} capitulos ya estaban traducidos, se omiten.")

    total = len(pendientes)
    if not pendientes:
        log("No hay nada pendiente por traducir.")
        progreso(total, total)
        return

    translator = GoogleTranslator(
        source=cfg.get("idioma_origen", "en"),
        target=cfg.get("idioma_destino", "es"),
    )

    for i, (archivo, destino) in enumerate(pendientes, start=1):
        if should_stop():
            log("Cancelado por el usuario.")
            return
        log(f"[{i}/{total}] Traduciendo: {archivo.name}")
        try:
            _traducir_docx(translator, archivo, destino, log)
        except Exception as e:
            log(f"[ERROR] Fallo '{archivo}': {e}")
        progreso(i, total)

    log("Traduccion completada.")
